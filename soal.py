import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os

# --- 1. KONFIGURASI API ---
def init_api():
    api_key = st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not api_key:
        st.error("⚠️ API Key tidak ditemukan!")
        st.stop()
    genai.configure(api_key=api_key)
    try:
        models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        target = next((m for m in models if "1.5-flash" in m), models[0])
        return genai.GenerativeModel(target)
    except:
        return genai.GenerativeModel('gemini-1.5-flash')

model = init_api()

# --- 2. FUNGSI EKSPOR WORD (LOGIKA TABEL DIPERKUAT) ---
def export_to_word(text, school_name, mapel):
    doc = Document()
    # Judul Identitas
    h = doc.add_heading('APLIKASI - SMPN 2 KALIPARE', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Satuan Pendidikan: {school_name}\nMata Pelajaran: {mapel}\nKelas / Fase: VII / Fase D\nKurikulum: Kurikulum Merdeka\n" + "-"*40)
    
    table_data = []
    in_table = False
    
    lines = text.split('\n')
    for line in lines:
        clean = line.strip()
        # Deteksi Baris Tabel Markdown
        if '|' in clean:
            # Abaikan baris pemisah markdown ---|---
            if all(c in '|- : ' for c in clean) and clean != "":
                continue
            # Ambil data sel
            cells = [c.strip() for c in clean.split('|') if c.strip() or clean.startswith('|')]
            if cells:
                table_data.append(cells)
                in_table = True
        else:
            # Jika tabel berakhir, buat tabel di Word
            if in_table and table_data:
                try:
                    num_cols = max(len(r) for r in table_data)
                    table = doc.add_table(rows=len(table_data), cols=num_cols)
                    table.style = 'Table Grid'
                    for r, row in enumerate(table_data):
                        for c, val in enumerate(row):
                            if c < num_cols:
                                table.cell(r, c).text = val
                except:
                    pass
                doc.add_paragraph("") # Spasi setelah tabel
                table_data = []
                in_table = False
            
            # Tambahkan teks non-tabel
            if clean:
                if clean.startswith('###'):
                    doc.add_heading(clean.replace('###', ''), level=1)
                elif clean.startswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(clean.replace('**', '')).bold = True
                else:
                    doc.add_paragraph(clean)
                    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 3. UI STREAMLIT ---
st.set_page_config(page_title="GuruAI - SMPN 2 Kalipare", layout="wide")
st.markdown('<div style="background:#1e3c72;color:white;padding:20px;border-radius:10px;text-align:center;"><h1>SMP NEGERI 2 KALIPARE</h1><p>Generator Soal HOTS & Perangkat Administrasi</p></div>', unsafe_allow_html=True)

col1, col2 = st.columns([2, 1])

with col1:
    source = st.radio("Input Materi:", ["Teks Manual", "Upload PDF"])
    materi_text = ""
    if source == "Teks Manual":
        materi_text = st.text_area("Tempelkan Materi Pelajaran:", height=300)
    else:
        f = st.file_uploader("Pilih PDF Materi", type=["pdf"])
        if f:
            pdf = PdfReader(f)
            for p in pdf.pages: materi_text += p.extract_text()

with col2:
    st.write("⚙️ **Pengaturan Komposisi Soal**")
    mapel_name = st.text_input("Mata Pelajaran", "Seni Rupa")
    
    bentuk_list = st.multiselect(
        "Pilih Bentuk Soal:", 
        ["Pilihan Ganda", "PG Kompleks", "Menjodohkan", "Isian Singkat", "Uraian"],
        default=["Pilihan Ganda", "Menjodohkan"]
    )
    
    dict_jumlah = {}
    if bentuk_list:
        for b in bentuk_list:
            dict_jumlah[b] = st.number_input(f"Jumlah {b}:", min_value=0, max_value=15, value=1)
    
    if st.button("Generate Perangkat Lengkap ✨"):
        if materi_text:
            with st.spinner("Sedang menyusun soal HOTS dan tabel kartu..."):
                try:
                    rincian = ", ".join([f"{v} soal {k}" for k, v in dict_jumlah.items() if v > 0])
                    prompt = (
                        f"Instansi: SMP NEGERI 2 KALIPARE. Mapel: {mapel_name}. Materi: {materi_text[:2500]}.\n"
                        f"Tugas: Buat soal HOTS (Level L3) dengan rincian: {rincian}.\n\n"
                        f"STRUKTUR WAJIB (HARUS TABEL MARKDOWN):\n"
                        f"1. ### KISI-KISI SOAL: Buat satu tabel (No|TP|Materi|Indikator|Level|No Soal).\n"
                        f"2. ### KARTU SOAL: Setiap nomor soal WAJIB dibuat dalam tabel Markdown terpisah.\n"
                        f"   Baris Tabel: | Komponen | Deskripsi |\n"
                        f"   Isi: CP, TP, Materi Utama, Buku Sumber, Indikator, Level (L3), Bentuk, No/Kunci/Skor, Stimulus, Rumusan Soal, Opsi.\n"
                        f"3. ### NASKAH SOAL: \n"
                        f"   - Khusus Menjodohkan: WAJIB buat tabel 2 kolom (| Pernyataan A | Pilihan B |).\n"
                        f"   - Bentuk lain: Teks rapi.\n"
                        f"4. ### KUNCI JAWABAN."
                    )
                    
                    res = model.generate_content(prompt)
                    output = res.text
                    st.markdown(output)
                    
                    st.download_button(
                        "📥 Download Word", 
                        export_to_word(output, "SMP NEGERI 2 KALIPARE", mapel_name), 
                        f"Perangkat_{mapel_name}.docx",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Kesalahan: {e}")
